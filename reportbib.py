# -*- coding: utf-8 -*-
"""
Created on Fri Aug 19 19:04:05 2022

@author: lan
"""

import pandas as pd
from docx import Document #pip install python-docx 

def df_to_doc(doc, df):
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    t.style = "TableGrid"

    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    return doc

def results_to_excel(ba, f_name="report.xlsx", index=0, exclude=[], 
                     autofit=True, M_len=100):
    #def to_excel(self, index=0, exclude=[], autofit=True, M_len=100):
    writer = pd.ExcelWriter(ba.res_folder + "\\reports\\report.xlsx", engine="xlsxwriter")
        
    #exclude += ["prepocess_df", "prod_df_plt"]
        
    for att in dir(ba):
        if ("_df" in att) and (att not in exclude) and ("_tex" not in att) and ("_ind" not in att):
            df = getattr(ba, att)
            try:
                df.to_excel(writer, sheet_name=att.split("_df")[0], index=index)
            except:
                print("Problem with", att)
                continue
            if autofit:
                for idx, col in enumerate(df):
                    series = df[col]
                    max_len = max((series.astype(str).map(len).max(),  
                                   len(str(series.name)))) + 1
                    if M_len is not None:
                        max_len = min(max_len, M_len)
                    writer.sheets[att.split("_df")[0]].set_column(idx, idx, max_len)
    writer.close()
    
def results_to_word(ba, f_name="report.docx", include_non_core=False):
    doc = Document()
    doc.add_heading("Report", 0) # kasneje boš dal ba.ldf čez vse
                
    # Dataset
    doc.add_heading("Dataset", 1) 
    if hasattr(ba, "main_info_df"):
        doc.add_heading("Main info", 2)
        doc = df_to_doc(doc, ba.main_info_df) 
    if hasattr(ba, "doc_type_df"):
        doc.add_heading("Documents per type", 2)
        doc = df_to_doc(doc, ba.doc_type_df)
        im = [i for i in ba.images if i.name == "documents per type"]
        if len(im):
            #doc.add_heading(im.heading, 1) # preveri, zakaj ne deluje
            doc.add_picture(im[0].path + ".png")
    if hasattr(ba, "top_gcd_df"):
        doc.add_heading("Global cited documents", 2)
        doc = df_to_doc(doc, ba.top_gcd_df)
              
    if hasattr(ba, "production_df"):
        doc.add_heading("Scientific production", 2)
        im = [i for i in ba.images if i.name == "scientific production"]
        if len(im):
            doc.add_picture(im[0].path + ".png")
            
    # Sources     
    doc.add_heading("Sources", 1) 
    if hasattr(ba, "sources_df"):
        doc.add_heading("Top sources", 2)
        im = [i for i in ba.images if i.name == "top sources"]
        if len(im):
            doc.add_picture(im[0].path + ".png")
        # tukaj je bilo to, ne vem čemu služi: add_plot(ba.images, "sources n", doc) 
         

    if hasattr(ba, "authors_df"):
        doc.add_heading("Top authors", 2)
        im = [i for i in ba.images if i.name == "top authors"]
        if len(im):
            doc.add_picture(im[0].path + ".png")

    if hasattr(ba, "ca_countries_df"):
        doc.add_heading("Top coutries (of corresponding author)", 2)
        im = [i for i in ba.images if i.name == "top CA countries"]
        if len(im):
            doc.add_picture(im[0].path + ".png")
            
    if hasattr(ba, "country_collab_df"):
        doc.add_heading("Country collaboration production", 3)
        im = [i for i in ba.images if i.name == "country collaboration production"]
        if len(im):
            doc.add_picture(im[0].path + ".png")
            
    doc.save(ba.res_folder + "\\reports\\report.docx")
    

def results_to_latex(ba, f_name="report.tex", template_tex="template report.tex", split_string="TO BE HERE: ("):
    with open(template_tex, "rt") as f_in:
        cnt = f_in.read()

    r = [rr.split(")")[0] for rr in cnt.split(split_string)[1:]]
    images = [im.path.split("plots\\")[1]+".png" for im in ba.images] + [im.path.split("plots\\")[1]+".pdf" for im in ba.images]
    print(r)
    print(images)
    
    for rr in r:
        try:
            cnt = cnt.replace(split_string+rr+")", getattr(ba, rr))
        except:
            if rr in images:
                try:
                    cnt =cnt.replace(split_string+rr+")", rr)
                except:
                    cnt =cnt.replace(split_string+rr+")", "empty.png")
            else:
                cnt =cnt.replace(split_string+rr+")", "empty.png")
      
    with open(f_name, "wt") as f_out:
        f_out.write(cnt)
        